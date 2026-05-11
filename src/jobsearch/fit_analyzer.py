import re
from dataclasses import dataclass, field
from pathlib import Path

from jobsearch.config import FIT_THRESHOLD, SKILL_WEIGHTS
from jobsearch.profiles import (
    OPTIONAL_MARKERS,
    REQUIRED_MARKERS,
    SKILL_CANONICAL,
    SKILL_CONTEXT_WORDS,
    SKILL_DICTIONARY,
    match_skills_against_profile,
    normalize,
)

# ---------------------------------------------------------------------------
# JD Parsing
# ---------------------------------------------------------------------------


def read_jd(path_or_text: str) -> str:
    """Read JD text from .docx, .pdf, or raw string."""
    p = Path(path_or_text)
    if not p.exists():
        # Treat as raw text
        return path_or_text

    suffix = p.suffix.lower()
    if suffix == ".docx":
        try:
            import docx

            doc = docx.Document(str(p))
            parts = [para.text for para in doc.paragraphs if para.text.strip()]
            # Also read table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())
            return "\n".join(parts)
        except Exception as e:
            raise RuntimeError(f"Cannot read .docx: {e}")

    if suffix == ".pdf":
        try:
            import pdfplumber

            with pdfplumber.open(str(p)) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception as e:
            raise RuntimeError(f"Cannot read .pdf: {e}")

    if suffix in (".txt", ".md"):
        return p.read_text(encoding="utf-8")

    raise ValueError(f"Unsupported file type: {suffix}")


# ---------------------------------------------------------------------------
# JD signal extraction (v2 — feeds the modality, salary, sector dimensions)
# ---------------------------------------------------------------------------

_CUR_MAP = {
    "$": "USD",
    "USD": "USD",
    "EUR": "EUR",
    "€": "EUR",
    "COP": "COP",
    "MXN": "MXN",
    "ARS": "ARS",
    "CLP": "CLP",
    "PEN": "PEN",
}

_SECTOR_VOCAB = {
    "humanitarian",
    "humanitario",
    "ngo",
    "ong",
    "nonprofit",
    "non-profit",
    "saas",
    "b2b",
    "b2c",
    "fintech",
    "edtech",
    "adtech",
    "healthtech",
    "ecommerce",
    "e-commerce",
    "marketplace",
    "consulting",
    "agency",
    "research",
    "academia",
    "public",
    "government",
    "gobierno",
    "data platforms",
    "ai infrastructure",
    "machine learning",
    "logistics",
    "supply chain",
    "manufacturing",
    "real estate",
    "energy",
    "climate",
    "sustainability",
    "media",
    "gaming",
}

_SALARY_NUMBER = re.compile(
    r"""
    (?P<cur_pre>\$|usd|eur|€|cop|mxn|ars|clp|pen)?\s*
    (?P<low>\d{1,3}(?:[.,]\d{3})*(?:[.,]\d+)?|\d+(?:[.,]\d+)?)\s*
    (?P<k>k)?
    (?:\s*[-a]\s*
        (?P<high>\d{1,3}(?:[.,]\d{3})*(?:[.,]\d+)?|\d+(?:[.,]\d+)?)\s*
        (?P<high_k>k)?
    )?
    \s*(?P<cur_post>\$|usd|eur|€|cop|mxn|ars|clp|pen)?
    """,
    re.IGNORECASE | re.VERBOSE,
)


def _parse_amount(raw: str | None, has_k_suffix: bool) -> float | None:
    """Parse a numeric string with comma or dot as thousands separator."""
    if not raw:
        return None
    # Detect a real decimal portion: a single dot/comma with <=2 trailing digits
    # AND no k suffix (because "80.5k" would be 80,500, not 80.5).
    decimal_portion = None
    for sep in (".", ","):
        if raw.count(sep) == 1 and len(raw.split(sep)[-1]) <= 2 and not has_k_suffix:
            integer = raw.split(sep)[0].replace(",", "").replace(".", "")
            decimal_portion = raw.split(sep)[-1]
            cleaned = f"{integer}.{decimal_portion}"
            break
    if decimal_portion is None:
        cleaned = raw.replace(",", "").replace(".", "")
    try:
        value = float(cleaned)
    except ValueError:
        return None
    if has_k_suffix:
        value *= 1000
    return value


def extract_jd_signals(jd_text: str) -> dict:
    """Best-effort extraction of structured signals from raw JD text.

    Returns a dict with optional keys ``modality``, ``hybrid_days``,
    ``salary_min``, ``salary_max``, ``currency``, ``sector_tokens``. Keys are
    omitted when the JD does not contain that information; ``score_fit``
    then drops the matching dimension and redistributes its weight.

    Patterns cover English and Spanish JDs; non-LATAM JDs may under-extract
    and ``score_fit`` handles that gracefully.
    """
    signals: dict = {}
    text_norm = normalize(jd_text or "")

    # Modality. ``hybrid`` always wins because a hybrid JD usually mentions
    # both "hybrid" and "in-office N days" — we must not misread the second
    # half as onsite. Then "no remote" / "onsite" / "presencial" beats a
    # bare "remote" mention. Bare "remote" matches last.
    if re.search(r"\b(hybrid|hibrido|hybrid[\s-]?remote|hybrid[\s-]?model)\b", text_norm):
        signals["modality"] = "hybrid"
        days_match = re.search(
            r"(\d+)\s*(?:days?|d[ií]as?)\s*(?:per\s*week|a la semana|in[\s-]?office|onsite|presencial)",
            text_norm,
        )
        if days_match:
            try:
                signals["hybrid_days"] = int(days_match.group(1))
            except ValueError:
                pass
    elif re.search(r"\b(on[\s-]?site|in[\s-]?office|presencial|no remote)\b", text_norm):
        signals["modality"] = "onsite"
    elif re.search(
        r"\b(remote|fully[\s-]?remote|100%?\s*remote|trabajo\s*remoto|remoto)\b", text_norm
    ):
        signals["modality"] = "remote"

    # Salary — look near keywords first so we do not pull a random number.
    # Allow dots inside the captured window (salary uses ``60.000`` style):
    # only newline terminates the search context.
    salary_context = re.search(
        r"(?:salary|salario|compensation|compensaci[oó]n|pay|remuneration|sueldo|wage)[^\n]{0,80}",
        text_norm,
    )
    if salary_context:
        m = _SALARY_NUMBER.search(salary_context.group(0))
        if m:
            cur_token = (m.group("cur_pre") or m.group("cur_post") or "").upper()
            currency = _CUR_MAP.get(cur_token, "")
            low = _parse_amount(m.group("low"), bool(m.group("k")))
            high = (
                _parse_amount(m.group("high"), bool(m.group("high_k"))) if m.group("high") else None
            )
            if low is not None:
                signals["salary_min"] = low
                signals["salary_max"] = high if high is not None else low
                if currency:
                    signals["currency"] = currency

    # Sector tokens — multi-word phrases checked first to avoid double counting.
    multi_word = {term for term in _SECTOR_VOCAB if " " in term and term in text_norm}
    single_word = {
        term
        for term in _SECTOR_VOCAB
        if " " not in term and re.search(rf"\b{re.escape(term)}\b", text_norm)
    }
    found = multi_word | single_word
    if found:
        signals["sector_tokens"] = found

    return signals


def find_jd_in_folder(folder: Path) -> Path | None:
    """Auto-detect JD file in a company folder."""
    from jobsearch.config import JD_FILENAMES

    for name in JD_FILENAMES:
        candidate = folder / name
        if candidate.exists():
            return candidate
    # Fallback: any .docx not containing "CV" or "Carta"
    for f in folder.glob("*.docx"):
        if not any(w in f.stem.upper() for w in ("CV", "CARTA", "COVER", "RESUME")):
            return f
    for f in folder.glob("*.pdf"):
        if not any(w in f.stem.upper() for w in ("CV", "CARTA", "COVER", "RESUME")):
            return f
    return None


# ---------------------------------------------------------------------------
# Skill Extraction
# ---------------------------------------------------------------------------

YEARS_PATTERNS = [
    # Note: patterns run against ``normalize(text)`` which strips accents,
    # so ``años`` becomes ``anos``. Patterns must match the normalized form.
    r"(\d+)\+?\s*(?:anos?|years?)\s+(?:de\s+)?(?:experiencia|experience)",
    r"(?:al menos|minimum|minimo|min\.?)\s+(\d+)\s*(?:anos?|years?)",
    r"(\d+)\s*-\s*\d+\s*(?:anos?|years?)",  # range: take minimum
    r"experiencia\s+de\s+(\d+)\s*(?:anos?|years?)",
]


def extract_years_from_text(text: str) -> float:
    """Extract minimum years of experience required from JD text."""
    text_norm = normalize(text)
    years_found = []
    for pattern in YEARS_PATTERNS:
        matches = re.findall(pattern, text_norm)
        for m in matches:
            try:
                years_found.append(float(m))
            except ValueError:
                pass
    return min(years_found) if years_found else 0.0


def is_near_context(sentence: str) -> bool:
    """Check if a skill mention is near a context word (reduces false positives)."""
    s = normalize(sentence)
    return any(ctx in s for ctx in SKILL_CONTEXT_WORDS)


def classify_requirement(sentence: str) -> str:
    """Determine if a skill is required or optional based on surrounding text."""
    s = normalize(sentence)
    for marker in OPTIONAL_MARKERS:
        if marker in s:
            return "optional"
    for marker in REQUIRED_MARKERS:
        if marker in s:
            return "required"
    return "required"  # default: treat as required


def extract_skills(jd_text: str) -> list[dict]:
    """Extract skills from JD text with category, required flag, and years."""
    text_norm = normalize(jd_text)
    sentences = re.split(r"[.\n;]", jd_text)

    found: dict[str, dict] = {}  # normalized_skill → skill dict

    all_terms = []
    for category, terms in SKILL_DICTIONARY.items():
        for term in terms:
            all_terms.append((normalize(term), term, category))

    # Sort by length descending (match longer terms first, e.g., "power bi" before "bi")
    all_terms.sort(key=lambda x: len(x[0]), reverse=True)

    # Low-confidence single-word terms that need context to be valid
    low_confidence_terms = {
        "r",
        "word",
        "excel",
        "teams",
        "zoom",
        "slack",
        "desarrollo",
        "evaluacion",
        "informes",
        "colaboracion",
        "sem",
        "spanish",
        "espanol",
    }

    for norm_term, original_term, category in all_terms:
        if norm_term not in text_norm:
            continue
        if norm_term in found:
            continue  # already found a longer match covering this term

        # Find which sentence contains this term
        for sentence in sentences:
            if norm_term in normalize(sentence):
                if norm_term in low_confidence_terms and not is_near_context(sentence):
                    continue

                is_required = classify_requirement(sentence)
                years = extract_years_from_text(sentence) or 0.0

                found[norm_term] = {
                    "skill": SKILL_CANONICAL.get(norm_term, original_term),
                    "category": category,
                    "required": 1 if is_required == "required" else 0,
                    "years_required": years if years > 0 else None,
                    "matched": 0,
                    "proficiency_gap": None,
                }
                break

    return list(found.values())


# ---------------------------------------------------------------------------
# Fit Scoring
# ---------------------------------------------------------------------------


@dataclass
class FitReport:
    company: str
    role: str
    fit_score: float
    skills_score: float
    experience_score: float
    total_required: int
    matched_required: int
    total_optional: int
    matched_optional: int
    gaps: list[dict] = field(default_factory=list)
    matched_skills: list[dict] = field(default_factory=list)
    years_required: float = 0.0
    years_profile: float = 0.0
    recommendation: str = ""
    threshold: float = FIT_THRESHOLD
    # v2 dimension scores. ``None`` means the dimension was not measurable for
    # this JD/profile combination and its weight was redistributed across the
    # dimensions actually scored.
    modality_score: float | None = None
    salary_score: float | None = None
    sector_score: float | None = None
    dimensions_used: list[str] = field(default_factory=list)
    weights_applied: dict[str, float] = field(default_factory=dict)

    def __post_init__(self):
        self.recommendation = "APPLY ✓" if self.fit_score >= self.threshold else "REVIEW GAPS FIRST"

    @property
    def percentage(self) -> float:
        return round(self.fit_score * 100, 1)

    @property
    def skills_pct(self) -> float:
        return round(self.skills_score * 100, 1)

    @property
    def exp_pct(self) -> float:
        return round(self.experience_score * 100, 1)


def _score_modality(
    user_pref: str | None,
    jd_modality: str | None,
    jd_hybrid_days: int | None,
    user_hybrid_max: int | None,
) -> float | None:
    """Return 1.0 / 0.5 / 0.0 / None for the modality dimension.

    - ``None`` means we cannot score (either side silent).
    - ``1.0`` for an exact match (remote==remote, onsite==onsite, hybrid with
      JD days <= user max).
    - ``0.5`` for "close enough" (e.g. user accepts hybrid but JD is remote;
      user prefers remote but JD is hybrid).
    - ``0.0`` for incompatible (remote vs onsite, or hybrid days > user max).
    """
    if not user_pref or not jd_modality:
        return None
    user = user_pref.lower()
    jd = jd_modality.lower()
    if user == jd:
        if jd == "hybrid" and jd_hybrid_days is not None and user_hybrid_max is not None:
            return 1.0 if jd_hybrid_days <= user_hybrid_max else 0.0
        return 1.0
    # Adjacent-but-not-equal modalities get partial credit.
    if {user, jd} == {"remote", "hybrid"}:
        return 0.5
    if {user, jd} == {"hybrid", "onsite"}:
        return 0.5
    # remote vs onsite — incompatible.
    return 0.0


def _score_salary(
    jd_min: float | None,
    jd_max: float | None,
    jd_currency: str | None,
    user_floor_amount: float | None,
    user_floor_currency: str | None,
) -> float | None:
    """Return 1.0 / 0.5 / 0.0 / None for the salary dimension.

    Currency mismatch with no FX conversion path returns ``None`` (we do not
    invent exchange rates). When currencies match:

    - ``1.0`` if JD max ≥ user floor
    - ``0.5`` if JD max is within 20% below the floor
    - ``0.0`` if JD max is below 80% of the floor
    """
    if user_floor_amount is None or jd_max is None:
        return None
    if user_floor_currency and jd_currency and user_floor_currency != jd_currency:
        return None
    if jd_max >= user_floor_amount:
        return 1.0
    ratio = jd_max / user_floor_amount
    if ratio >= 0.8:
        return 0.5
    return 0.0


def _score_sector(
    jd_tokens: set[str] | None, user_targets: list[str] | None, user_avoids: list[str] | None
) -> float | None:
    """Jaccard-style score between JD sector tokens and user's targets.

    - Returns ``None`` if either side is empty.
    - Penalty: if the JD tokens overlap the user's ``sectors_avoid`` list,
      the score drops by 0.3 (floored at 0.0).
    """
    if not jd_tokens or not user_targets:
        return None
    targets_norm = {normalize(t) for t in user_targets}
    avoids_norm = {normalize(t) for t in (user_avoids or [])}
    overlap = jd_tokens & targets_norm
    union = jd_tokens | targets_norm
    base = len(overlap) / len(union) if union else 0.0
    if jd_tokens & avoids_norm:
        base = max(0.0, base - 0.3)
    return round(base, 4)


def score_fit(
    jd_skills: list[dict],
    profile: dict,
    *,
    company: str = "",
    role: str = "",
    jd_text: str = "",
    jd_signals: dict | None = None,
    user_filters: dict | None = None,
) -> FitReport:
    """Calculate the fit score across user-weighted dimensions.

    Dimensions:
      - ``skills`` and ``experience`` come from ``jd_skills`` + ``profile``
        (v1 behavior).
      - ``modality``, ``salary_floor``, ``sector_fit`` come from
        ``jd_signals`` (parsed once with ``extract_jd_signals``) and
        ``user_filters`` (loaded from ``roles/<role>.md`` and/or
        ``_brain/USER_CONTEXT.md``).

    A dimension that cannot be scored for this JD returns ``None`` and its
    weight is redistributed proportionally across the dimensions actually
    scored. ``FitReport.dimensions_used`` records which dimensions
    contributed and ``FitReport.weights_applied`` shows the final
    normalization, so the score is always explainable.
    """
    from jobsearch.config import scoring_threshold, scoring_weights

    required = [s for s in jd_skills if s.get("required", 1)]
    optional = [s for s in jd_skills if not s.get("required", 1)]

    matched_req = [s for s in required if s.get("matched")]
    matched_opt = [s for s in optional if s.get("matched")]
    gaps = [s for s in required if not s.get("matched")]

    # Skills score.
    if len(required) == 0:
        skills_score = 0.5
    else:
        base = len(matched_req) / len(required)
        opt_bonus = (len(matched_opt) / max(len(optional), 1)) * 0.1
        skills_score = min(1.0, base + opt_bonus)

    # Experience score.
    years_required = extract_years_from_text(jd_text) if jd_text else 0.0
    years_profile = float(profile.get("total_years_experience", 0))
    if years_required == 0:
        exp_score = 1.0
    else:
        shortfall = years_required - years_profile
        if shortfall <= 0:
            exp_score = 1.0
        elif shortfall <= 2:
            exp_score = 0.75
        elif shortfall <= 4:
            exp_score = 0.4
        else:
            exp_score = 0.1

    # New dimensions (v2) — optional, may be None.
    signals = jd_signals or {}
    filters = user_filters or {}

    modality_score = _score_modality(
        user_pref=filters.get("work_modality_preference"),
        jd_modality=signals.get("modality"),
        jd_hybrid_days=signals.get("hybrid_days"),
        user_hybrid_max=filters.get("hybrid_days_max"),
    )

    user_floor = filters.get("salary_floor") or {}
    salary_score = _score_salary(
        jd_min=signals.get("salary_min"),
        jd_max=signals.get("salary_max"),
        jd_currency=signals.get("currency"),
        user_floor_amount=user_floor.get("monthly"),
        user_floor_currency=user_floor.get("currency"),
    )

    sector_score = _score_sector(
        jd_tokens=signals.get("sector_tokens"),
        user_targets=filters.get("sectors_target"),
        user_avoids=filters.get("sectors_avoid"),
    )

    # Pick weights from config, drop dimensions that returned None, normalize.
    raw_weights = scoring_weights()
    dimension_values: dict[str, float] = {
        "skills": skills_score,
        "experience": exp_score,
    }
    if modality_score is not None:
        dimension_values["modality"] = modality_score
    if salary_score is not None:
        dimension_values["salary_floor"] = salary_score
    if sector_score is not None:
        dimension_values["sector_fit"] = sector_score

    # Keep only weights for dimensions that actually scored.
    applied = {dim: raw_weights[dim] for dim in dimension_values if dim in raw_weights}
    total = sum(applied.values())
    if total > 0:
        applied = {k: v / total for k, v in applied.items()}
    else:
        # Edge case: user disabled both skills and experience (unlikely, but
        # do not crash). Fall back to even split across remaining dims.
        applied = {k: 1 / len(dimension_values) for k in dimension_values}

    fit_score = sum(dimension_values[d] * w for d, w in applied.items())

    return FitReport(
        company=company,
        role=role,
        fit_score=round(fit_score, 4),
        skills_score=round(skills_score, 4),
        experience_score=round(exp_score, 4),
        total_required=len(required),
        matched_required=len(matched_req),
        total_optional=len(optional),
        matched_optional=len(matched_opt),
        gaps=gaps,
        matched_skills=matched_req,
        years_required=years_required,
        years_profile=years_profile,
        threshold=scoring_threshold(),
        modality_score=modality_score,
        salary_score=salary_score,
        sector_score=sector_score,
        dimensions_used=list(applied.keys()),
        weights_applied={k: round(v, 4) for k, v in applied.items()},
    )


# ---------------------------------------------------------------------------
# Report Generator
# ---------------------------------------------------------------------------


def generate_fit_report_md(report: FitReport, jd_source: str = "") -> str:
    from datetime import date

    # Build the dimensions table dynamically so we never show ``None`` rows.
    dim_rows = [
        f"| Skills match | {report.matched_required}/{report.total_required} ({report.skills_pct}%) |",
        f"| Experience | {report.years_profile} yrs vs {report.years_required} required ({report.exp_pct}%) |",
    ]
    if report.modality_score is not None:
        dim_rows.append(f"| Modality | {round(report.modality_score * 100, 1)}% |")
    if report.salary_score is not None:
        dim_rows.append(f"| Salary | {round(report.salary_score * 100, 1)}% |")
    if report.sector_score is not None:
        dim_rows.append(f"| Sector fit | {round(report.sector_score * 100, 1)}% |")
    dim_rows.append(f"| **TOTAL** | **{report.percentage}%** |")

    lines = [
        f"# Fit Report — {report.company}",
        f"**Role:** {report.role}  |  **Date:** {date.today().isoformat()}",
        f"**JD:** {jd_source}" if jd_source else "",
        "",
        "---",
        "",
        f"## Score: {report.percentage}% — {report.recommendation}",
        "",
        "| Dimension | Score |",
        "|-----------|-------|",
        *dim_rows,
        "",
        "---",
        "",
        "## Matched Skills",
    ]

    if report.matched_skills:
        for s in report.matched_skills:
            lines.append(f"- ✅ {s['skill']} ({s.get('category', '')})")
    else:
        lines.append("_No required skill from the JD was found in your profile._")

    lines += [
        "",
        "## Gaps (required by the JD, missing from your profile)",
    ]

    if report.gaps:
        for g in report.gaps:
            yr = f" — {g['years_required']} years required" if g.get("years_required") else ""
            lines.append(f"- ❌ {g['skill']} ({g.get('category', '')}){yr}")
    else:
        lines.append("_No gaps detected._")

    if report.total_optional > 0:
        lines += [
            "",
            f"## Optional skills ({report.matched_optional}/{report.total_optional} present)",
        ]

    # Weights footer — only show the dimensions that actually contributed so
    # the user sees the math used for this specific report.
    if report.weights_applied:
        lines += [
            "",
            "---",
            "",
            "## Weights applied (normalized)",
            "",
            "| Dimension | Weight |",
            "|-----------|--------|",
            *(f"| {dim} | {round(w * 100, 1)}% |" for dim, w in report.weights_applied.items()),
        ]

    lines += [
        "",
        "---",
        f"_Threshold to apply: {int(report.threshold * 100)}%_",
    ]

    return "\n".join(line for line in lines)


def analyze(jd_path_or_text: str, role: str, company: str = "", save_to: Path = None) -> FitReport:
    """Full pipeline: read JD → extract skills → match → score → report.

    v2 wires up the modality / salary / sector dimensions automatically by
    pulling ``jd_signals`` from ``extract_jd_signals`` and ``user_filters``
    from ``profiles.load_role_filters(role)``. Both are optional — if a
    signal is missing or the user has not filled their filters, the
    corresponding dimension is skipped and its weight is redistributed.
    """
    from jobsearch.profiles import load_profile, load_role_filters

    jd_text = read_jd(jd_path_or_text)
    profile = load_profile(role)
    skills = extract_skills(jd_text)
    skills = match_skills_against_profile(skills, profile)

    jd_signals = extract_jd_signals(jd_text)
    user_filters = load_role_filters(role)

    report = score_fit(
        skills,
        profile,
        company=company,
        role=role,
        jd_text=jd_text,
        jd_signals=jd_signals,
        user_filters=user_filters,
    )

    if save_to:
        md = generate_fit_report_md(report, jd_source=str(jd_path_or_text))
        save_to.write_text(md, encoding="utf-8")

    return report, skills
